#!/usr/bin/env python3
"""
Genera MEMORIA_TFC.docx desde MEMORIA_TFC.md
Formato según guía del centro: Arial 12, interlineado 1.5, márgenes académicos
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

MD_PATH = "/Users/alvaromartinez/Github/PRICEWISE/docs/MEMORIA_TFC.md"
DOCX_PATH = "/Users/alvaromartinez/Github/PRICEWISE/docs/MEMORIA_TFC.docx"

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GRAY_BG = RGBColor(0xF5, 0xF5, 0xF5)
DARK_GRAY = RGBColor(0x44, 0x44, 0x44)


def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3) + Cm(1)   # 3cm + 1cm encuadernación
        section.right_margin = Cm(3)


def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def apply_paragraph_format(para, size_pt=12, bold=False, color=None, space_before=0, space_after=6, align=None):
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if align:
        pf.alignment = align
    for run in para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color


def add_heading(doc, text, level):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.bold = True
    run.font.color.rgb = NAVY

    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    if level == 1:
        run.font.size = Pt(16)
        pf.space_before = Pt(18)
        pf.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(14)
        pf.space_before = Pt(14)
        pf.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(12)
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)
    else:
        run.font.size = Pt(11)
        pf.space_before = Pt(8)
        pf.space_after = Pt(3)

    return para


def add_body(doc, text):
    """Añade párrafo de cuerpo con formato base. Soporta negrita/código inline."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Parsear negrita (**texto**) y código (`texto`)
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            run = para.add_run(part)

        run.font.name = 'Arial'
        if not (part.startswith('`') and part.endswith('`')):
            run.font.size = Pt(12)

    return para


def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.left_indent = Cm(0.5 + level * 0.5)

    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            run = para.add_run(part)

        run.font.name = 'Arial'
        if not (part.startswith('`') and part.endswith('`')):
            run.font.size = Pt(12)

    return para


def add_code_block(doc, lines):
    code_text = '\n'.join(lines)
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.left_indent = Cm(0.5)
    pf.right_indent = Cm(0.5)

    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # Fondo gris para el bloque de código
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    para._p.get_or_add_pPr().append(shd)

    return para


def parse_table_line(line):
    """Extrae celdas de una línea de tabla markdown."""
    cells = [c.strip() for c in line.strip().strip('|').split('|')]
    return cells


def add_table(doc, rows):
    if not rows:
        return
    # Filtrar separador (---|---|---)
    data = [r for r in rows if not re.match(r'^[\s|:\-]+$', r)]
    if not data:
        return

    parsed = [parse_table_line(r) for r in data]
    cols = max(len(r) for r in parsed)
    # Normalizar número de columnas
    parsed = [r + [''] * (cols - len(r)) for r in parsed]

    table = doc.add_table(rows=len(parsed), cols=cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(parsed):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            para = cell.paragraphs[0]
            pf = para.paragraph_format
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)

            pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
            parts = pattern.split(cell_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = para.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                else:
                    run = para.add_run(part)

                run.font.name = 'Arial'
                run.font.size = Pt(10)

            if i == 0:
                for run in para.runs:
                    run.bold = True
                set_cell_background(cell, '1B2A4A')
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif i % 2 == 0:
                set_cell_background(cell, 'EEF2FF')

    doc.add_paragraph()


def add_page_break(doc):
    para = doc.add_paragraph()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    para._p.append(br)


def process_md(doc, md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code = False
    code_lines = []
    table_rows = []
    in_table = False

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Bloque de código
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
                i += 1
                continue
            else:
                in_code = False
                add_code_block(doc, code_lines)
                code_lines = []
                i += 1
                continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Tabla markdown
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(line)
            i += 1
            continue
        else:
            if in_table:
                in_table = False
                add_table(doc, table_rows)
                table_rows = []

        # Separador ---
        if re.match(r'^-{3,}\s*$', line):
            i += 1
            continue

        # Headings
        if line.startswith('#### '):
            add_heading(doc, line[5:].strip(), 4)
        elif line.startswith('### '):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith('## '):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith('# '):
            # Salto de página antes de cada sección principal
            if doc.paragraphs:
                add_page_break(doc)
            add_heading(doc, line[2:].strip(), 1)

        # Bullets
        elif re.match(r'^(\s*)[-*] ', line):
            indent = len(re.match(r'^(\s*)', line).group(1)) // 2
            text = re.sub(r'^(\s*)[-*] ', '', line)
            add_bullet(doc, text, indent)

        # Línea en blanco
        elif line.strip() == '':
            pass  # no añadir párrafos vacíos extra

        # Blockquote (> texto)
        elif line.startswith('> '):
            para = doc.add_paragraph()
            pf = para.paragraph_format
            pf.left_indent = Cm(1)
            pf.space_before = Pt(4)
            pf.space_after = Pt(4)
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            run = para.add_run(line[2:].strip())
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.color.rgb = DARK_GRAY

        # Párrafo normal
        else:
            if line.strip():
                add_body(doc, line)

        i += 1

    # Flush tabla si quedó pendiente
    if in_table and table_rows:
        add_table(doc, table_rows)


def main():
    doc = Document()
    set_margins(doc)

    # Estilo base del documento
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    process_md(doc, MD_PATH)

    doc.save(DOCX_PATH)
    print(f"✓ Generado: {DOCX_PATH}")


if __name__ == '__main__':
    import docx
    main()
